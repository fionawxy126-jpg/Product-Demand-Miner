"""
报告生成模块 — 按模板输出 Markdown + Word (.docx) + CSV 报告
"""

import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _generate_summary(product_name, analysis, product_scope=""):
    """根据分析结果生成简明摘要"""
    meta = analysis["meta"]
    pain_points = analysis["pain_points"]
    features = analysis["features"]

    parts = [f"本次调研共抓取 {meta['total_posts']} 条帖子，"]

    if pain_points:
        top_pains = [pp["label"] for pp in pain_points[:3]]
        parts.append(f"识别出 {len(pain_points)} 类用户痛点，其中「{'、'.join(top_pains)}」最为突出。")

    if features:
        top_feats = [f["label"] for f in features[:2]]
        parts.append(f"用户最期待的功能包括「{'、'.join(top_feats)}」。")

    scope_hint = f"（{product_scope}方向）" if product_scope else ""
    parts.append(f"建议 {product_name}{scope_hint}团队重点关注上述痛点，在产品迭代中优先解决高频问题。")

    return "".join(parts) + "（请人工补充修改）"


def _pain_section(pain_points):
    lines = []
    for i, pp in enumerate(pain_points, 1):
        top = pp["posts"][0]
        lines.append(f"### {i}. {pp['label']}")
        lines.append("")
        lines.append(f"- **帖子**: {top['title']}")
        lines.append(f"  {top['score']} 赞 | {top['num_comments']} 评论 | {top.get('subreddit', '')} | {top['created_utc']}")
        lines.append(f"  链接: {top['url']}")
        lines.append(f"- **相关帖子数量**: {pp['count']} 条")
        lines.append(f"- **核心问题**: {_infer_core_issue(top)}")
        lines.append("")

        if len(pp["posts"]) > 1:
            lines.append("  其他相关帖子:")
            for p in pp["posts"][1:4]:
                lines.append(f"  - [{p['title']}]({p['url']}) — {p['score']} 赞 | {p.get('subreddit', '')} | {p['created_utc']}")
            lines.append("")

    return "\n".join(lines)


def _infer_core_issue(post):
    """从帖子内容推断核心问题描述"""
    text = (post["title"] + " " + post.get("body", "")[:200]).lower()

    issues = [
        (["prompt", "instruction", "error on", "misunderstood"],
         "用户指令不够精确，AI 生成大量代码后发现方向错误，需要推倒重来"),
        (["billing", "charged", "cost me", "hidden fee", "hermes", "silently"],
         "存在隐性扣费或计费不透明的问题，用户在不知情的情况下被额外收费"),
        (["token", "ccusage", "visibility"],
         "用户无法清晰了解 token 消耗的具体去向和分布，缺乏透明度"),
        (["pro plan", "paywall", "locked", "no longer", "price"],
         "产品订阅计划突然变更，核心功能被移到更高价位，用户感到被欺骗"),
        (["limit", "quota", "rate limit", "allowance", "throttl"],
         "订阅额度不足以支撑正常使用频率，用户被迫降低使用量或寻找替代方案"),
        (["loop", "stuck", "infinite", "overnight", "burning", "spiraled"],
         "AI Agent 进入死循环或失控状态，持续消耗资源无法自行停止"),
        (["memory", "context", "forget", "claude.md", "amnesia"],
         "AI 在长会话中遗忘之前约定的上下文或指令，导致行为不一致"),
        (["hallucinat", "made up", "fabricat", "nonexistent"],
         "AI 编造不存在的 API、库或事实，导致生成的代码无法运行"),
        (["data loss", "lost code", "deleted", "destroyed", "wiped", "nuked"],
         "AI 操作导致用户代码或数据被意外删除或覆盖，造成不可逆损失"),
        (["slow", "performance", "latency", "hang", "timeout"],
         "工具响应速度慢或频繁超时，严重影响开发效率"),
        (["crash", "unstable", "downtime", "outage", "bug"],
         "工具频繁崩溃或出现稳定性问题，无法可靠地完成工作"),
        (["privacy", "security", "data safe", "leak", "sensitive"],
         "用户担心代码和数据的隐私安全，对数据传输和存储方式缺乏信任"),
    ]

    for keywords, desc in issues:
        if any(k in text for k in keywords):
            return desc

    return "用户在使用过程中遇到了影响体验的关键问题"


def _feature_section(features):
    lines = []
    for feat in features:
        lines.append(f"### {feat['label']}")
        lines.append("")
        for p in feat["posts"][:3]:
            lines.append(f"- [{p['title']}]({p['url']}) — {p['score']} 赞 | {p.get('subreddit', '')} | {p['created_utc']}")
        lines.append("")
    return "\n".join(lines)


def _ranking_table(pain_points):
    lines = [
        "| 排名 | 痛点类别 | 相关帖子数 | 最高赞数 | 核心诉求 |",
        "|:----:|---------|:---------:|:-------:|---------|",
    ]
    for i, pp in enumerate(pain_points, 1):
        top = pp["posts"][0]
        core = _infer_core_issue(top)
        if len(core) > 40:
            core = core[:40] + "..."
        lines.append(f"| {i} | {pp['label']} | {pp['count']} | {pp['top_score']} | {core} |")
    return "\n".join(lines)


def generate_report(product_name, keywords, subreddits, time_range,
                    analysis, platforms=None, product_desc="", product_scope="",
                    output_dir="output"):
    """
    生成 Markdown 报告 + CSV。

    Args:
        product_name:  产品名称
        keywords:      关键词列表
        subreddits:    搜索的板块列表
        time_range:    时间范围描述
        analysis:      analyze_posts() 返回值
        platforms:     数据来源平台列表
        output_dir:    输出目录

    Returns:
        (md_path, csv_path)
    """
    os.makedirs(output_dir, exist_ok=True)

    meta = analysis["meta"]
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    safe_name = product_name.replace(" ", "_").replace("/", "_")
    base = f"{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

    sources = ", ".join(platforms) if platforms else "Reddit"
    subs = ", ".join(meta["subreddits"])

    desc_line = f"\n> 产品简介: {product_desc}" if product_desc else ""
    scope_line = f"\n> 产品范围: {product_scope}" if product_scope else ""

    summary = _generate_summary(product_name, analysis, product_scope)

    md = f"""# {product_name} 开发者痛点与市场调研报告

> 生成时间: {now}
> 数据来源: {sources}
> 共抓取 **{meta['total_posts']}** 条帖子{desc_line}{scope_line}

## 搜索范围

- **板块**: {subs}
- **关键词**: {', '.join(keywords)}
- **时间范围**: {time_range}
- **总评论数**: {meta['total_comments']:,} 条
- **平均点赞数**: {meta['avg_score']}

---

## 摘要

{summary}

---

## 一、用户痛点

{_pain_section(analysis['pain_points'])}
## 二、用户期待的功能

{_feature_section(analysis['features'])}
## 三、痛点数量排序

{_ranking_table(analysis['pain_points'])}

---

## 四、热门帖子 Top 10

| # | 标题 | 板块 | 点赞 | 评论 | 日期 |
|:--:|------|------|:----:|:----:|:----:|
"""
    for i, p in enumerate(analysis["top10"], 1):
        md += f"| {i} | [{p['title']}]({p['url']}) | {p.get('subreddit', '')} | {p['score']:,} | {p['num_comments']} | {p['created_utc']} |\n"

    md += f"""
---

## 五、总结

[待人工审核填写：4-5 条核心洞察]

---

*报告由 Product Demand Miner 自动生成*
"""

    md_path = os.path.join(output_dir, f"{base}.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md)

    # CSV
    csv_path = os.path.join(output_dir, f"{base}.csv")
    keys = ["id", "title", "body", "subreddit", "score", "num_comments", "url", "created_utc", "keyword", "source"]
    seen = set()
    rows = []
    for p in analysis["top10"] + [p for pp in analysis["pain_points"] for p in pp["posts"]]:
        if p["id"] not in seen:
            seen.add(p["id"])
            rows.append(p)

    with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=keys, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)

    return md_path, csv_path


def generate_docx(product_name, keywords, subreddits, time_range,
                  analysis, platforms=None, product_desc="", product_scope="",
                  output_dir="output"):
    """生成可编辑的 Word (.docx) 报告"""
    os.makedirs(output_dir, exist_ok=True)

    meta = analysis["meta"]
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    safe_name = product_name.replace(" ", "_").replace("/", "_")
    base = f"{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    sources = ", ".join(platforms) if platforms else "Reddit"

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)

    # 标题
    title = doc.add_heading(f'{product_name} 开发者痛点与市场调研报告', level=0)

    # 元信息
    meta_para = doc.add_paragraph()
    meta_para.add_run(f'生成时间: {now}').italic = True
    meta_para.add_run('\n')
    run = meta_para.add_run(f'数据来源: {sources}')
    run.italic = True
    meta_para.add_run('\n')
    run = meta_para.add_run(f'共抓取 {meta["total_posts"]} 条帖子')
    run.italic = True
    run.bold = True
    if product_desc:
        meta_para.add_run('\n')
        run = meta_para.add_run(f'产品简介: {product_desc}')
        run.italic = True
    if product_scope:
        meta_para.add_run('\n')
        run = meta_para.add_run(f'产品范围: {product_scope}')
        run.italic = True

    # 搜索范围
    doc.add_heading('搜索范围', level=2)
    doc.add_paragraph(f'板块: {", ".join(meta["subreddits"])}')
    doc.add_paragraph(f'关键词: {", ".join(keywords)}')
    doc.add_paragraph(f'时间范围: {time_range}')
    doc.add_paragraph(f'总评论数: {meta["total_comments"]:,} 条')
    doc.add_paragraph(f'平均点赞数: {meta["avg_score"]}')

    # 摘要
    doc.add_heading('摘要', level=2)
    summary = _generate_summary(product_name, analysis, product_scope)
    doc.add_paragraph(summary)

    # 一、用户痛点
    doc.add_heading('一、用户痛点', level=1)
    for i, pp in enumerate(analysis["pain_points"], 1):
        top = pp["posts"][0]
        doc.add_heading(f'{i}. {pp["label"]}', level=3)

        p = doc.add_paragraph()
        p.add_run('帖子: ').bold = True
        p.add_run(top['title'])
        doc.add_paragraph(f'{top["score"]} 赞 | {top["num_comments"]} 评论 | {top.get("subreddit", "")} | {top["created_utc"]}')

        link = doc.add_paragraph()
        link.add_run('链接: ')
        link.add_run(top['url'])

        p = doc.add_paragraph()
        p.add_run('相关帖子数量: ').bold = True
        p.add_run(f'{pp["count"]} 条')

        p = doc.add_paragraph()
        p.add_run('核心问题: ').bold = True
        p.add_run(_infer_core_issue(top))

        if len(pp["posts"]) > 1:
            doc.add_paragraph('其他相关帖子:')
            for post in pp["posts"][1:4]:
                doc.add_paragraph(
                    f'{post["title"]} — {post["score"]} 赞 | {post.get("subreddit", "")} | {post["created_utc"]}',
                    style='List Bullet',
                )
        doc.add_paragraph()

    # 二、用户期待的功能
    doc.add_heading('二、用户期待的功能', level=1)
    for feat in analysis["features"]:
        doc.add_heading(feat["label"], level=3)
        for p in feat["posts"][:3]:
            doc.add_paragraph(
                f'{p["title"]} — {p["score"]} 赞 | {p.get("subreddit", "")} | {p["created_utc"]}',
                style='List Bullet',
            )
        doc.add_paragraph()

    # 三、痛点数量排序
    doc.add_heading('三、痛点数量排序', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Shading Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = '排名'
    hdr[1].text = '痛点类别'
    hdr[2].text = '帖子数'
    hdr[3].text = '最高赞'
    hdr[4].text = '核心诉求'

    for i, pp in enumerate(analysis["pain_points"], 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = pp["label"]
        row[2].text = str(pp["count"])
        row[3].text = str(pp["top_score"])
        core = _infer_core_issue(pp["posts"][0])
        row[4].text = core[:60] + '...' if len(core) > 60 else core

    # 四、热门帖子 Top 10
    doc.add_heading('四、热门帖子 Top 10', level=1)
    t2 = doc.add_table(rows=1, cols=5)
    t2.style = 'Light Shading Accent 1'
    h2 = t2.rows[0].cells
    h2[0].text = '#'
    h2[1].text = '标题'
    h2[2].text = '板块'
    h2[3].text = '赞'
    h2[4].text = '评论'

    for i, p in enumerate(analysis["top10"], 1):
        row = t2.add_row().cells
        row[0].text = str(i)
        row[1].text = p["title"]
        row[2].text = p.get("subreddit", "")
        row[3].text = str(p["score"])
        row[4].text = str(p["num_comments"])

    # 五、总结（占位）
    doc.add_heading('五、总结', level=1)
    placeholder2 = doc.add_paragraph('[待人工审核填写：4-5 条核心洞察]')
    placeholder2.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()
    footer = doc.add_paragraph('报告由 Product Demand Miner 自动生成')
    footer.runs[0].italic = True
    footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    docx_path = os.path.join(output_dir, f"{base}.docx")
    doc.save(docx_path)

    return docx_path
